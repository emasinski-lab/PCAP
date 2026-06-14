#!/usr/bin/env python3
"""
Script pour convertir le guide Markdown en document Word (.docx)

Usage:
    python scripts/convert_to_word.py
    python scripts/convert_to_word.py GUIDE_UTILISATION.md PCAP_Guide.docx
"""

import sys
import os
import re
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def install_docx():
    """Installe python-docx si nécessaire"""
    import subprocess
    print("Installation de python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    print("python-docx installé avec succès!")


def add_formatted_text(paragraph, text):
    """Ajoute du texte avec mise en forme (gras, italique, liens)"""
    # Pattern pour les liens [texte](url)
    link_pattern = r'\[(.*?)\]\((.*?)\)'
    
    # Traiter les liens d'abord
    parts = re.split(link_pattern, text)
    for i, part in enumerate(parts):
        if i % 3 == 0:  # Texte normal
            if part:
                # Traiter le gras et l'italique dans le texte normal
                add_formatted_text_recursive(paragraph, part)
        elif i % 3 == 1:  # Texte du lien
            if part:
                run = paragraph.add_run(part)
                run.font.color.rgb = RGBColor(0, 0, 255)  # Bleu
                run.font.underline = True
        # i % 3 == 2 est l'URL, on l'ignore pour l'affichage


def add_formatted_text_recursive(paragraph, text):
    """Ajoute du texte avec mise en forme récursive"""
    # Pattern pour le gras **texte**
    bold_matches = list(re.finditer(r'\*\*(.*?)\*\*', text))
    # Pattern pour l'italique *texte*
    italic_matches = list(re.finditer(r'\*(.*?)\*', text))
    
    if not bold_matches and not italic_matches:
        # Pas de mise en forme, ajouter le texte tel quel
        paragraph.add_run(text)
        return
    
    # Trouver la première occurrence de mise en forme
    all_matches = []
    for match in bold_matches:
        all_matches.append((match.start(), match.end(), 'bold', match.group(1)))
    for match in italic_matches:
        # Éviter les conflits avec le gras
        if not any(match.start() >= b[0] and match.end() <= b[1] for b in bold_matches):
            all_matches.append((match.start(), match.end(), 'italic', match.group(1)))
    
    if not all_matches:
        paragraph.add_run(text)
        return
    
    # Trier par position
    all_matches.sort(key=lambda x: x[0])
    
    # Ajouter le texte avant la première mise en forme
    first_match = all_matches[0]
    if first_match[0] > 0:
        paragraph.add_run(text[:first_match[0]])
    
    # Ajouter les parties avec mise en forme
    for i, (start, end, style_type, content) in enumerate(all_matches):
        # Ajouter le texte entre la fin de la précédente et le début de celle-ci
        if i > 0:
            prev_end = all_matches[i-1][1]
            if start > prev_end:
                paragraph.add_run(text[prev_end:start])
        
        # Ajouter le texte avec mise en forme
        run = paragraph.add_run(content)
        if style_type == 'bold':
            run.bold = True
        elif style_type == 'italic':
            run.italic = True
        
        # Si c'est la dernière mise en forme, ajouter le texte restant
        if i == len(all_matches) - 1 and end < len(text):
            paragraph.add_run(text[end:])
    
    # Ajouter le texte après la dernière mise en forme
    last_match = all_matches[-1]
    if last_match[1] < len(text):
        paragraph.add_run(text[last_match[1]:])


def add_table(doc, rows):
    """Ajoute un tableau au document"""
    if not rows:
        return
    
    # Nettoyer les lignes vides
    rows = [row for row in rows if row.strip()]
    if not rows:
        return
    
    # Créer le tableau
    table = doc.add_table(rows=len(rows), cols=1)
    table.style = 'Table Grid'
    
    # Remplir le tableau
    for i, row in enumerate(rows):
        # Diviser la ligne par |
        cells = [cell.strip() for cell in row.split('|') if cell.strip()]
        
        # Ajouter les cellules
        for j, cell in enumerate(cells):
            if j >= len(table.rows[i].cells):
                # Ajouter une colonne si nécessaire
                table.rows[i].add_cell()
            table.rows[i].cells[j].text = cell


def markdown_to_word(md_file, docx_file=None):
    """Convertit un fichier Markdown en document Word"""
    
    if not DOCX_AVAILABLE:
        install_docx()
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml.ns import qn
    
    # Si aucun fichier de sortie spécifié, utiliser un nom par défaut
    if docx_file is None:
        base_name = os.path.splitext(os.path.basename(md_file))[0]
        docx_file = f"{base_name}.docx"
    
    print(f"Conversion de {md_file} vers {docx_file}...")
    
    # Lire le fichier Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Créer le document Word
    doc = Document()
    
    # Configurer les styles
    styles = doc.styles
    
    # Style pour les titres
    heading1 = styles['Heading 1']
    heading1.font.name = 'Calibri'
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 0, 139)  # Bleu foncé
    
    heading2 = styles['Heading 2']
    heading2.font.name = 'Calibri'
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(70, 130, 180)  # Bleu steel
    
    heading3 = styles['Heading 3']
    heading3.font.name = 'Calibri'
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.font.color.rgb = RGBColor(0, 100, 0)  # Vert foncé
    
    # Style pour le texte normal
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    
    # Style pour le code
    code_style = styles.add_style('CodeStyle', 1)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(10)
    code_style.font.color.rgb = RGBColor(128, 0, 0)  # Rouge foncé
    
    # Style pour les tableaux
    table_style = styles.add_style('TableStyle', 1)
    table_style.font.name = 'Calibri'
    table_style.font.size = Pt(10)
    
    # Traiter le contenu Markdown
    lines = md_content.split('\n')
    in_code_block = False
    in_table = False
    table_rows = []
    
    for line in lines:
        line = line.strip()
        
        # Ignorer les lignes vides
        if not line:
            if not in_code_block and not in_table:
                doc.add_paragraph()
            continue
        
        # Détecter les titres
        if line.startswith('# '):
            if in_code_block:
                in_code_block = False
                doc.add_paragraph()
            if in_table:
                add_table(doc, table_rows)
                in_table = False
                table_rows = []
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            if in_code_block:
                in_code_block = False
                doc.add_paragraph()
            if in_table:
                add_table(doc, table_rows)
                in_table = False
                table_rows = []
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            if in_code_block:
                in_code_block = False
                doc.add_paragraph()
            if in_table:
                add_table(doc, table_rows)
                in_table = False
                table_rows = []
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            if in_code_block:
                in_code_block = False
                doc.add_paragraph()
            if in_table:
                add_table(doc, table_rows)
                in_table = False
                table_rows = []
            doc.add_heading(line[5:], level=4)
        
        # Détecter les séparateurs
        elif line.startswith('---') or line.startswith('***') or line.startswith('___'):
            if in_code_block:
                in_code_block = False
                doc.add_paragraph()
            if in_table:
                add_table(doc, table_rows)
                in_table = False
                table_rows = []
            # Ajouter une ligne horizontale
            p = doc.add_paragraph()
            p.add_run('─' * 50).font.color.rgb = RGBColor(200, 200, 200)
        
        # Détecter les blocs de code
        elif line.startswith('```'):
            if in_code_block:
                # Fin du bloc de code
                in_code_block = False
            else:
                # Début du bloc de code
                in_code_block = True
                if in_table:
                    add_table(doc, table_rows)
                    in_table = False
                    table_rows = []
        
        # Détecter les tableaux
        elif '|' in line and not in_code_block:
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(line)
        
        # Détecter les listes
        elif line.startswith('- ') or line.startswith('* ') or line.startswith('+ '):
            if in_code_block:
                in_code_block = False
                doc.add_paragraph()
            if in_table:
                add_table(doc, table_rows)
                in_table = False
                table_rows = []
            # Ajouter un élément de liste
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(line[2:].strip())
        
        elif re.match(r'^\d+\.', line):
            if in_code_block:
                in_code_block = False
                doc.add_paragraph()
            if in_table:
                add_table(doc, table_rows)
                in_table = False
                table_rows = []
            # Ajouter un élément de liste numérotée
            p = doc.add_paragraph(style='List Number')
            p.add_run(line.split('.', 1)[1].strip())
        
        # Détecter les liens
        elif line.startswith('|') and in_table:
            table_rows.append(line)
        
        # Texte normal ou code inline
        else:
            if in_code_block:
                # Ajouter le code
                p = doc.add_paragraph(style='CodeStyle')
                p.add_run(line)
            elif in_table:
                table_rows.append(line)
            else:
                # Texte normal
                p = doc.add_paragraph()
                add_formatted_text(p, line)
    
    # Fermer les blocs ouverts
    if in_code_block:
        in_code_block = False
    if in_table:
        add_table(doc, table_rows)
    
    # Ajouter un pied de page
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run(f"Généré le {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").font.size = Pt(8)
    p.add_run(" - PCAP Analyzer v1.0").font.size = Pt(8)
    
    # Sauvegarder le document
    doc.save(docx_file)
    print(f"✓ Document Word sauvegardé: {docx_file}")
    return docx_file


def main():
    if len(sys.argv) < 2:
        # Utiliser le guide par défaut
        md_file = "GUIDE_UTILISATION.md"
        docx_file = "PCAP_Guide_Utilisation.docx"
    else:
        md_file = sys.argv[1]
        docx_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    if not os.path.exists(md_file):
        print(f"Erreur: Le fichier {md_file} n'existe pas")
        return 1
    
    try:
        markdown_to_word(md_file, docx_file)
        return 0
    except Exception as e:
        print(f"Erreur lors de la conversion: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == '__main__':
    sys.exit(main())
